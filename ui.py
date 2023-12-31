"""
Streamlit UI for Pypercraft
"""

import os
import streamlit as st
from docx import Document
from pypercraft.pypercraft import Pypercraft

st.set_page_config(page_title="Pypercraft")


def main():
    """
    Runs an example of Pypercraft
    """

    st.title("Pypercraft")
    st.subheader("Generate and Download Quality Papers with GenAI")

    html_content = '''
    <div style="display: flex; flex-direction: row; align-items: center;">
        <a href="https://github.com/alkhalifas/pypercraft" target="_blank">
            <img src="https://badgen.net/badge/icon/GitHub?icon=github&label" alt="Repo">
        </a>
        <a style="margin-left: 10px;" href="https://pypi.org/project/pypercraft/" target="_blank">
            <img src="https://badgen.net/pypi/v/pypercraft" alt="Repo">
        </a>
        <a style="margin-left: 10px;" href="https://github.com/alkhalifas/pypercraft" target="_blank">
            <img src="https://badgen.net/github/license/micromatch/micromatch" alt="Repo">
        </a>
    </div>
    '''

    # Display the HTML content using st.write()
    st.write(html_content, unsafe_allow_html=True)
    st.markdown("<br>", unsafe_allow_html=True)


    # User input
    query = st.text_input("Enter your query:", placeholder="Enter your query here...")
    topic = st.text_input("Enter the topic:", placeholder="Enter the topic here...")

    # Separate columns
    col1, col2 = st.columns(2)

    with col1:
        num_pages = st.number_input("Number of pages:", min_value=1, value=2, step=1, max_value=3)
    with col2:
        tone = st.selectbox("Select Tone", ["Professional", "Cute", "Artistic", "Moderate"])

    # Get API Key
    api_key = st.text_input("Enter your API Key:", placeholder="xxxx-xxxxx-xxxx-xxxxx")

    # Generate paper on button click
    if st.button("Generate Paper"):
        if api_key:
            try:
                pypercraft = Pypercraft(query, topic, num_pages, tone, api_key)
            except ValueError:
                st.error(' Invalid API Key. Can you double check your input?')
        else:
            api_key = os.getenv("OPENAI_API_KEY")
            pypercraft = Pypercraft(query, topic, num_pages, tone, api_key)

        with st.spinner("Generating Paper..."):
            paper = pypercraft.construct()
            st.session_state.paper = paper

        st.subheader("Title")
        st.write(paper["title"])

        st.subheader("Introduction")
        st.write(paper["introduction"])

        st.subheader("Body")
        st.write(paper["body"])

        st.subheader("Conclusion")
        st.write(paper["conclusion"])
        export_document(st.session_state.paper)


def export_document(paper):
    """
    Generates and exports the docx document
    """
    doc = Document()

    # Adding title, introduction, body, and conclusion to the DOCX document
    doc.add_heading(paper["title"], level=1)
    doc.add_paragraph(paper["introduction"])
    doc.add_paragraph(paper["body"])
    doc.add_paragraph(paper["conclusion"])

    # Save as DOCX
    doc_filename = "tmp/generated_paper.docx"
    doc.save(doc_filename)
    st.success(f"DOCX file generated: {doc_filename}")

    # Read the content of the temporary file
    with open(doc_filename, "rb") as file:
        data = file.read()

    # Display a download button with the file content
    st.download_button(
        label="Download docx",
        data=data,
        file_name="pypercraft_paper.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


if __name__ == "__main__":
    main()
